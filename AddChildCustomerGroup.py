## Test case for adding a new child to a existent costumer group ##
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from docx import Document
from docx.shared import Inches
from pagemethods.customergroups import count_customers

import os
import datetime

document = Document()


## Initialize chrome
driver = webdriver.Chrome()
driver.get("https://cinq.repairq.io/")


## Login
elem = driver.find_element_by_id("UserLoginForm_username")
elem.send_keys("pedro")
elem = driver.find_element_by_id("UserLoginForm_password")
elem.send_keys("pedro")
elem.send_keys(Keys.RETURN)

## Identify number of Customer Groups
driver.get("https://cinq.repairq.io/customerGroups")
count = count_customers(driver)
old_count = len(driver.find_elements_by_class_name("largest-row"))
print('count', count)

## Add new Customer Group
driver.get("https://cinq.repairq.io/customerGroups/add")
    # fill Customer Groups fields
elem = driver.find_element_by_id("CustomerGroup_name")
elem.send_keys("School number " + str(count+1))
elem = driver.find_element_by_id("CustomerGroup_description")
elem.send_keys("This is description for School number " + str(count+1) + ", added for testing.")

    # select parent_
select = Select(driver.find_element_by_id("CustomerGroup_parent_id"))
select.select_by_index(count-3)
selected = select.first_selected_option.text
print('selected', selected)

### Screenshot of selected parent
driver.fullscreen_window()
driver.implicitly_wait(5)
driver.save_screenshot('selected_parent.png')
driver.maximize_window()
document.add_picture('selected_parent.png', width=Inches(6.25))


elem = driver.find_element_by_id("save-button")
#elem.click()   ####CHANGE HERE
driver.get("https://cinq.repairq.io/customerGroups/15")

driver.implicitly_wait(5)

### Screenshot of summary page, after save
driver.fullscreen_window()
driver.implicitly_wait(5)
driver.save_screenshot('created_customer_group.png')
driver.maximize_window()
document.add_picture('created_customer_group.png', width=Inches(6.25))

d = datetime.datetime.now()
filename = str(d.year) + str(d.month) + str(d.day) + str(d.hour) + str(d.minute)+ str(d.second)+ ".docx"

elem = driver.find_elements_by_class_name("section-header")[1]
if elem.text == "Parent":
    elem = driver.find_elements_by_class_name("controls")[4]
    if elem.text == selected:
        filename = "Success - Add child customer group " + filename
        document.save(filename)
        print("test ok")
        print("Parent on Summary:  " + elem.text)
        print("Parent selected:  " + selected)
    else:
        filename = "Failed - Add child customer group" + filename
        document.save(filename)
        print("test nao ok")
        print("Parent on Summary:  " + elem.text)
        print("Parent selected:  " + selected)
else:
    print("no parent")


os.remove("created_customer_group.png")
os.remove("selected_parent.png")
driver.quit()


